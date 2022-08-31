from docx import Document
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH

job = "Werkstudent"
job_field = "Strukturdynamik"
my_adress = "Gräfstr. 121A, 81241 München"
phone = "+49 176 20674272"
city = "München, den "
email = "vincent.vincent@haw-hamburg.de"
company_name = "EDAG"
available_date = "01.10.2022"
available_duracy = "6 Monate"
application_title = f"BEWERBUNG ALS {job.upper()} - {job_field.upper()}"
opening_text = f"Sehr geehrte Damen und Herren,\nich möchte Ihnen gleich zu Beginn Gründe nennen, warum Sie von mir als neuem {job}en profitieren werden:"
soft_skills = [
    "Selbstständige, zuverlässige, und strukturierte Arbeitsweise und lösungsorientiert",
    "Rasche Auffassungsgabe",
    "Hohe Einsatzbereitschaft und Verantwortungsbewusstsein",
    "Freude an Zusammenarbeit im Team",
    "Lernbereitschaft"
]

motivation_list = [
    "Work-Life Balance",
    "Attraktive Vergütung",
    "Entwicklungsmöglichkeit",
    "Start-Up Arbeitskultur"
]
reason = "blabla"

closing_text = f"Gerne möchte ich meine bisherigen Erfahrungen aus dem Studium und Beruf in dieser Gelegenheit mitbringen, um einen Beitrag in Ihrem Unternehmen als ein Teil des Teams leisten zu können und ich bin zuversichtlich, dass ich mit Ihrem Unternehmen sogar mit noch mehr wertvollen Erfahrungen bereichert werden kann.\n\nIch freue mich auf eine persönliches Gespräch und überzeuge Sie gerne im persönlichen Gespräch. Ihrer positiven Rückmeldung sehe ich mit großer Freude entgegen. Außerdem stehe ich für die ausgeschriebene Stelle ab {available_date} für {available_duracy}.\n\nMit freundlichen Grüßen\nVincent"

#Personal Info paragraph
datum_unformatted = str(date.today())
datum_list = datum_unformatted.split("-")
datum = datum_list[2] + "." + datum_list[1] + "." + datum_list[0]

document = Document()

document.add_heading('Vincent', 0)

document.add_paragraph(my_adress + "\n" + phone + "\n" + email)

#Date paragraph
datum_para = document.add_paragraph(city + datum)
datum_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Company paragraph
company_para = document.add_paragraph()
company_para.add_run(company_name.upper()).bold = True

# application paragraph
application_para = document.add_paragraph()
application_para.add_run(application_title).bold = True

# Opening paragraph
opening_para = document.add_paragraph(opening_text)

# Skill list
for skill in soft_skills:
    document.add_paragraph(skill, style='List Bullet')

# Motivation
document.add_paragraph(f"Darum möchte ich gerne bei {company_name.upper()} als {job} im Bereich {job_field} sein:")
for motivation in motivation_list:
    document.add_paragraph(motivation, style="List Bullet")

# Closing paragraph
document.add_paragraph(closing_text)

document.save(company_name + "_" + job.lower() + "_" + job_field.lower().replace(" ", "_") + ".docx")